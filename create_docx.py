"""
Script to convert RiskMind Tech Stack to Word document
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    return table

def create_tech_stack_docx():
    doc = Document()
    
    # Title
    doc.add_heading('RiskMind - Project-Specific Tech Stack', 0)
    
    # Document Info
    doc.add_heading('Document Information', level=1)
    add_table(doc, ['Property', 'Value'], [
        ['Project', 'RiskMind Underwriting Co-Pilot'],
        ['Platform', 'Windows 10/11'],
        ['Date', 'February 2026'],
        ['Version', '1.1']
    ])
    doc.add_paragraph()
    
    # Runtime
    doc.add_heading('1. Runtime Requirements', level=1)
    add_table(doc, ['Software', 'Version', 'Download'], [
        ['Node.js', '20+', 'nodejs.org'],
        ['Python', '3.11+', 'python.org'],
        ['Git', 'Latest', 'git-scm.com']
    ])
    doc.add_paragraph()
    
    # Frontend
    doc.add_heading('2. Frontend Stack', level=1)
    add_table(doc, ['Package', 'Version', 'Purpose'], [
        ['react', '18.x', 'UI framework'],
        ['react-dom', '18.x', 'DOM rendering'],
        ['react-router-dom', '6.x', 'Page navigation'],
        ['axios', '1.x', 'API calls'],
        ['lucide-react', 'Latest', 'Icons'],
        ['vite', '5.x', 'Build tool'],
        ['typescript', '5.x', 'Type safety'],
        ['tailwindcss', '4.x', 'Styling']
    ])
    doc.add_paragraph()
    
    # Backend
    doc.add_heading('3. Backend Stack', level=1)
    add_table(doc, ['Package', 'Version', 'Purpose'], [
        ['fastapi', '0.109+', 'REST API framework'],
        ['uvicorn', '0.27+', 'Web server'],
        ['sqlalchemy', '2.x', 'Database ORM'],
        ['pydantic', '2.x', 'Data validation'],
        ['python-dotenv', '1.x', 'Environment config'],
        ['aiosqlite', '0.19+', 'SQLite driver (local)'],
        ['asyncpg', '0.29+', 'PostgreSQL driver (AWS)'],
        ['httpx', '0.26+', 'HTTP client'],
        ['openai', '1.x', 'OpenAI SDK'],
        ['boto3', '1.34+', 'AWS SDK (Bedrock)']
    ])
    doc.add_paragraph()
    
    # Database
    doc.add_heading('4. Database', level=1)
    add_table(doc, ['Environment', 'Database', 'Purpose'], [
        ['Local Development', 'SQLite', 'Easy setup'],
        ['AWS Production', 'PostgreSQL (RDS)', 'Scalable'],
        ['Vector Search', 'OpenSearch', 'RAG']
    ])
    doc.add_paragraph()
    
    # AWS Services
    doc.add_heading('5. AWS Services Required', level=1)
    add_table(doc, ['Service', 'Purpose', 'Required'], [
        ['Amazon Bedrock', 'LLM (Claude 3.5 Sonnet)', 'Yes'],
        ['Amazon RDS', 'PostgreSQL database', 'Yes'],
        ['Amazon OpenSearch', 'Vector search for RAG', 'Yes'],
        ['AWS App Runner', 'Backend hosting', 'Yes'],
        ['AWS Amplify', 'Frontend hosting', 'Yes'],
        ['Amazon S3', 'File storage', 'Yes'],
        ['Amazon Cognito', 'User authentication', 'Yes'],
        ['AWS Secrets Manager', 'API keys storage', 'Yes'],
        ['Amazon CloudWatch', 'Logs & monitoring', 'Yes'],
        ['AWS IAM', 'Access management', 'Yes']
    ])
    doc.add_paragraph()
    
    # Bedrock Models
    doc.add_heading('6. Bedrock Models to Request (Priority Order)', level=1)
    add_table(doc, ['Priority', 'Model', 'Provider', 'Purpose'], [
        ['HIGHEST', 'Claude 3.5 Sonnet', 'Anthropic', 'PRIMARY - Risk analysis'],
        ['HIGH', 'Claude 3 Haiku', 'Anthropic', 'Fast responses, chat'],
        ['HIGH', 'Titan Embeddings V2', 'Amazon', 'Vector search (RAG)'],
        ['MEDIUM', 'Claude 3 Opus', 'Anthropic', 'Complex analysis']
    ])
    doc.add_paragraph()
    
    # AWS Access Requirements
    doc.add_heading('7. AWS Access Requirements', level=1)
    
    doc.add_heading('Option A: Full Console Access (Recommended)', level=2)
    doc.add_paragraph('Developers get AWS Console login with ability to view and manage resources.')
    
    doc.add_heading('Option B: Programmatic Access Only (Minimum)', level=2)
    doc.add_paragraph('IT Admin sets up resources, developers receive credentials to use in code.')
    doc.add_paragraph()
    
    # Minimum Required
    doc.add_heading('MINIMUM ACCESS REQUIRED (Must Have)', level=2)
    doc.add_paragraph('If IT cannot provide full console access, developers MUST receive these credentials:')
    add_table(doc, ['Credential', 'Format', 'Purpose'], [
        ['AWS Access Key ID', 'AKIA...', 'API authentication'],
        ['AWS Secret Access Key', 'wJalr...', 'API authentication'],
        ['AWS Region', 'us-east-1', 'Service location'],
        ['RDS Connection String', 'postgresql://user:pass@host:5432/db', 'Database access'],
        ['OpenSearch Endpoint', 'https://search-xxx.us-east-1.es.amazonaws.com', 'RAG vector search'],
        ['Cognito User Pool ID', 'us-east-1_xxxxxx', 'User authentication'],
        ['Cognito Client ID', 'xxxxxxx', 'App authentication'],
        ['S3 Bucket Name', 'riskmind-documents', 'File storage']
    ])
    doc.add_paragraph()
    
    # Good to Have
    doc.add_heading('GOOD TO HAVE (Recommended)', level=2)
    doc.add_paragraph('These console permissions make development faster:')
    add_table(doc, ['Service', 'Access Level', 'Why It Helps'], [
        ['Amazon Bedrock', 'Full invoke', 'Test prompts in Playground'],
        ['Amazon CloudWatch', 'Read logs', 'Debug production issues'],
        ['AWS App Runner', 'View + Deploy', 'See deployment status'],
        ['AWS Amplify', 'View + Deploy', 'See frontend builds'],
        ['Amazon RDS', 'Read only', 'View database metrics'],
        ['Amazon S3', 'Read + Write', 'Browse uploaded files']
    ])
    doc.add_paragraph()
    
    # Action Items
    doc.add_heading('8. Action Items for IT Team', level=1)
    items = [
        '1. Enable Amazon Bedrock in region (us-east-1)',
        '2. Request access to Claude 3.5 Sonnet, Claude 3 Haiku, Titan Embeddings',
        '3. Create RDS PostgreSQL instance',
        '4. Create OpenSearch Service domain',
        '5. Create Cognito user pool',
        '6. Create S3 bucket for documents',
        '7. Create IAM users with permissions above',
        '8. Provide credentials to development team'
    ]
    for item in items:
        doc.add_paragraph(item)
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('9. Summary', level=1)
    add_table(doc, ['Category', 'Count', 'Details'], [
        ['Frontend packages', '8', 'React, Vite, Router, Axios, Tailwind, TypeScript, Lucide'],
        ['Backend packages', '10', 'FastAPI, SQLAlchemy, Pydantic, OpenAI, Boto3'],
        ['Databases', '3', 'SQLite (local), PostgreSQL (prod), OpenSearch (RAG)'],
        ['AWS Services', '10', 'Bedrock, RDS, OpenSearch, App Runner, Amplify, S3, Cognito, Secrets, CloudWatch, IAM'],
        ['Total', '~30', 'Core technologies only']
    ])
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph('Document Version: 1.1 | Last Updated: February 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    doc.save('docs/RiskMind-Tech-Stack.docx')
    print('Word document created: docs/RiskMind-Tech-Stack.docx')

if __name__ == '__main__':
    create_tech_stack_docx()
